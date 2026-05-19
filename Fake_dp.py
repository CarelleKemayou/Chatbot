import os
from docx import Document
from docx.shared import Pt, RGBColor
import openpyxl
import csv

BASE = os.path.join(r"C:\Users\Carelle.Kemayou\Desktop", "base_donnees")

# Créer tous les sous-dossiers
for folder in ["RH", "Finance", "Juridique", "Comptabilite", "IT", "Direction", "Marketing", "Achats"]:
    os.makedirs(os.path.join(BASE, folder), exist_ok=True)

def make_docx(path, title, sections):
    doc = Document()
    doc.add_heading(title, 0)
    for heading, content in sections:
        doc.add_heading(heading, level=1)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph(content)
    doc.save(path)
    print(f"✅ {path}")

def make_csv(path, headers, rows):
    with open(path, 'w', newline='', encoding='utf-8') as f:
        w = csv.writer(f)
        w.writerow(headers)
        w.writerows(rows)
    print(f"✅ {path}")

def make_xlsx(path, sheet_name, headers, rows):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet_name
    ws.append(headers)
    for row in rows:
        ws.append(row)
    wb.save(path)
    print(f"✅ {path}")

# ─── RH ───
make_docx(os.path.join(BASE, "RH", "reglement_interieur.docx"), "Règlement Intérieur — Mairie de Villeneuve-les-Prés", [
    ("Article 1 — Objet", "Le présent règlement intérieur s'applique à l'ensemble du personnel de la Mairie de Villeneuve-les-Prés, titulaires et contractuels, conformément à la loi du 26 janvier 1984 portant dispositions statutaires relatives à la fonction publique territoriale."),
    ("Article 2 — Horaires de travail", "La durée hebdomadaire de travail est fixée à 35 heures. Les horaires sont organisés comme suit :\n- Lundi au vendredi : 8h30 – 12h00 / 13h30 – 17h00\n- Permanences le samedi matin de 9h à 12h pour les services à la population."),
    ("Article 3 — Congés annuels", "Tout agent titulaire bénéficie de 25 jours ouvrés de congés annuels, auxquels s'ajoutent les jours de fractionnement selon les conditions définies par le décret n°85-1250 du 26 novembre 1985."),
    ("Article 4 — Discipline", "Tout manquement aux obligations professionnelles peut faire l'objet d'une sanction disciplinaire. Les sanctions sont prononcées par l'autorité territoriale après avis du Conseil de discipline."),
    ("Article 5 — Télétravail", "Le télétravail est autorisé jusqu'à 2 jours par semaine pour les postes éligibles, sous réserve de l'accord du responsable hiérarchique et de la signature d'une convention individuelle de télétravail."),
])

make_docx(os.path.join(BASE, "RH", "procedure_recrutement.docx"), "Procédure de Recrutement — Mairie de Villeneuve-les-Prés", [
    ("1. Identification du besoin", "Le responsable de service exprime le besoin de recrutement via le formulaire DRHE-01 transmis à la Direction des Ressources Humaines au moins 3 mois avant la date souhaitée de prise de poste."),
    ("2. Publication de l'offre", "L'offre est publiée sur la Bourse de l'Emploi Territorial (BET), le site de la mairie, et Pôle Emploi. La durée minimale de publication est de 4 semaines."),
    ("3. Sélection des candidatures", "Les candidatures sont examinées par un jury composé du DRH, du responsable de service concerné, et d'un élu référent. Un entretien de motivation est organisé pour les candidats présélectionnés."),
    ("4. Décision et intégration", "La décision de recrutement est notifiée par courrier recommandé. L'agent intègre la collectivité après signature du contrat ou arrêté de nomination, et suit un parcours d'intégration de 2 jours."),
])

make_csv(os.path.join(BASE, "RH", "effectifs_2025.csv"),
    ["Matricule", "Nom", "Prénom", "Service", "Grade", "Statut", "Date_entrée", "Salaire_brut"],
    [
        ["M001", "Durand", "Sophie", "Direction Générale", "Attaché territorial", "Titulaire", "2015-03-01", "3200"],
        ["M002", "Lefebvre", "Marc", "Finance", "Rédacteur principal", "Titulaire", "2018-06-15", "2750"],
        ["M003", "Moreau", "Isabelle", "RH", "Technicien territorial", "Titulaire", "2019-09-01", "2400"],
        ["M004", "Petit", "Antoine", "Urbanisme", "Ingénieur", "Titulaire", "2020-01-10", "3100"],
        ["M005", "Bernard", "Claire", "Communication", "Rédacteur", "Contractuel", "2022-04-01", "2200"],
        ["M006", "Thomas", "Lucas", "IT", "Technicien systèmes", "Contractuel", "2021-11-15", "2600"],
        ["M007", "Richard", "Émilie", "Achats", "Adjoint administratif", "Titulaire", "2017-02-28", "2100"],
        ["M008", "Simon", "Paul", "Juridique", "Attaché juridique", "Titulaire", "2016-07-01", "3050"],
        ["M009", "Laurent", "Nadia", "Comptabilité", "Rédacteur financier", "Titulaire", "2014-05-20", "2900"],
        ["M010", "Michel", "Théo", "Espaces verts", "Agent technique", "Titulaire", "2023-03-01", "2000"],
    ]
)

# ─── FINANCE ───
make_xlsx(os.path.join(BASE, "Finance", "budget_2025.xlsx"), "Budget 2025",
    ["Service", "Chapitre", "Libellé", "Budget_voté", "Engagé", "Disponible"],
    [
        ["Direction", "011", "Charges à caractère général", 120000, 45000, 75000],
        ["RH", "012", "Charges de personnel", 2500000, 1800000, 700000],
        ["Travaux", "020", "Dépenses imprévues", 80000, 12000, 68000],
        ["IT", "021", "Immobilisations corporelles", 95000, 60000, 35000],
        ["Communication", "65", "Autres charges de gestion courante", 35000, 18000, 17000],
        ["Achats", "60", "Achats et variations de stocks", 150000, 90000, 60000],
        ["Urbanisme", "16", "Emprunts et dettes", 300000, 300000, 0],
        ["Culture", "65", "Subventions associations", 75000, 40000, 35000],
        ["Sport", "65", "Subventions clubs sportifs", 50000, 30000, 20000],
        ["Environnement", "23", "Immobilisations en cours", 200000, 80000, 120000],
    ]
)

make_docx(os.path.join(BASE, "Finance", "rapport_financier_2024.docx"), "Rapport Financier Annuel 2024 — Mairie de Villeneuve-les-Prés", [
    ("Synthèse de l'exercice 2024", "L'exercice budgétaire 2024 s'est clôturé avec un excédent de fonctionnement de 187 432 euros, en amélioration de 12% par rapport à 2023. Le taux d'exécution budgétaire global atteint 91,4%."),
    ("Recettes de fonctionnement", "Les recettes totales s'élèvent à 8 245 000 euros, dont 4 100 000 euros de dotations de l'État (DGF), 2 800 000 euros de fiscalité locale, et 1 345 000 euros de recettes propres (services, redevances, subventions)."),
    ("Dépenses de fonctionnement", "Les dépenses de fonctionnement représentent 8 057 568 euros. Les charges de personnel constituent le premier poste avec 5 200 000 euros (64,5% du total). Les charges à caractère général s'élèvent à 1 850 000 euros."),
    ("Investissements réalisés", "Le programme d'investissement 2024 a été exécuté à hauteur de 76%. Les principaux projets concernent la rénovation énergétique de la salle polyvalente (340 000 €), le remplacement du parc informatique (95 000 €), et les travaux de voirie (210 000 €)."),
    ("Perspectives 2025", "Le budget primitif 2025 est arrêté à 8 900 000 euros en fonctionnement et 1 200 000 euros en investissement. Une attention particulière sera portée à la maîtrise des charges énergétiques et à la transition numérique des services."),
])

# ─── JURIDIQUE ───
make_docx(os.path.join(BASE, "Juridique", "marches_publics_en_cours.docx"), "Marchés Publics en Cours — 2025", [
    ("Marché n°2025-001 — Fournitures de bureau", "Objet : Fourniture de consommables et matériels de bureau pour l'ensemble des services.\nProcédure : Marché à procédure adaptée (MAPA)\nMontant estimé : 45 000 € HT\nTitulaire : OFFICE PLUS SAS\nDurée : 12 mois renouvelable 2 fois\nDate de notification : 15 janvier 2025"),
    ("Marché n°2025-002 — Maintenance informatique", "Objet : Maintenance préventive et corrective du parc informatique municipal (120 postes).\nProcédure : MAPA\nMontant estimé : 78 000 € HT / an\nTitulaire : TECHSOLUTIONS 95\nDurée : 24 mois\nDate de notification : 1er février 2025"),
    ("Marché n°2025-003 — Travaux de voirie", "Objet : Réfection de la rue des Lilas et de l'avenue du Général de Gaulle.\nProcédure : Appel d'offres ouvert\nMontant estimé : 420 000 € HT\nTitulaire : BTP ÎLE-DE-FRANCE\nDurée : 6 mois\nDate de notification : 10 mars 2025"),
    ("Marché n°2025-004 — Nettoyage des locaux", "Objet : Nettoyage quotidien de l'Hôtel de Ville, de la médiathèque et du gymnase municipal.\nProcédure : MAPA\nMontant estimé : 62 000 € HT / an\nTitulaire : PROPRE & NET SERVICES\nDurée : 12 mois renouvelable"),
])

make_docx(os.path.join(BASE, "Juridique", "conventions_partenaires.docx"), "Conventions de Partenariat Actives", [
    ("Convention avec le CCAS", "La Mairie et le Centre Communal d'Action Sociale (CCAS) ont signé le 1er janvier 2025 une convention de partenariat pour la mise à disposition de locaux et de personnels. La Mairie prend en charge 30% du budget de fonctionnement du CCAS, soit 95 000 euros pour l'exercice 2025."),
    ("Convention avec l'association Les Ailes du Sport", "Convention d'objectifs et de moyens signée le 15 février 2025. Subvention annuelle de 18 000 euros en contrepartie de l'organisation de 4 événements sportifs ouverts au public et de la gestion du terrain de football municipal."),
    ("Convention avec la Communauté de Communes", "Convention de mutualisation des services informatiques avec la Communauté de Communes du Val Vert. Partage des coûts à hauteur de 40% pour la Mairie de Villeneuve-les-Prés. Économies estimées : 25 000 euros/an."),
])

# ─── COMPTABILITE ───
make_xlsx(os.path.join(BASE, "Comptabilite", "suivi_depenses_T1_2025.xlsx"), "Dépenses T1 2025",
    ["Date", "N_mandat", "Fournisseur", "Objet", "Service", "Montant_HT", "TVA", "Montant_TTC", "Statut"],
    [
        ["2025-01-08", "M2025-001", "OFFICE PLUS SAS", "Fournitures de bureau Jan", "Direction", 1200, 240, 1440, "Payé"],
        ["2025-01-15", "M2025-002", "EDF PRO", "Électricité Hôtel de Ville", "Patrimoine", 3800, 760, 4560, "Payé"],
        ["2025-01-22", "M2025-003", "TECHSOLUTIONS 95", "Maintenance informatique Jan", "IT", 3250, 650, 3900, "Payé"],
        ["2025-02-05", "M2025-004", "PROPRE & NET", "Nettoyage locaux Jan-Fév", "Patrimoine", 5166, 1033, 6199, "Payé"],
        ["2025-02-12", "M2025-005", "IMPRIMERIE DUPONT", "Impression bulletins municipaux", "Communication", 2100, 420, 2520, "Payé"],
        ["2025-02-20", "M2025-006", "BTP ÎLE-DE-FRANCE", "Acompte travaux voirie", "Travaux", 84000, 16800, 100800, "Payé"],
        ["2025-03-03", "M2025-007", "FRANCE TELECOM PRO", "Téléphonie fixe T1", "IT", 980, 196, 1176, "Payé"],
        ["2025-03-10", "M2025-008", "ASSURANCES PUBLIQUES", "Prime assurance annuelle", "Direction", 12500, 0, 12500, "Payé"],
        ["2025-03-18", "M2025-009", "LIBRAIRIE MUNICIPALE", "Abonnements presse et revues", "Documentation", 450, 0, 450, "En cours"],
        ["2025-03-25", "M2025-010", "GARAGE MUNICIPAL", "Entretien véhicules de service", "Technique", 3200, 640, 3840, "En cours"],
    ]
)

# ─── IT ───
make_docx(os.path.join(BASE, "IT", "inventaire_parc_informatique.docx"), "Inventaire du Parc Informatique — Mairie de Villeneuve-les-Prés", [
    ("Synthèse du parc", "Au 1er janvier 2025, le parc informatique de la mairie comprend 120 postes de travail, 8 serveurs, 15 imprimantes réseau, et 25 équipements mobiles (tablettes et smartphones)."),
    ("Postes de travail", "- 85 PC fixes (Dell OptiPlex 7090, Windows 11 Pro) — âge moyen : 2,3 ans\n- 35 PC portables (HP EliteBook 840, Windows 11 Pro) — âge moyen : 1,8 ans\nLogiciels : Suite Microsoft 365, SEDIT (logiciel RH), CIVIL-NET (état civil), GEO-MAIRIE (SIG)"),
    ("Infrastructure serveurs", "- 2 serveurs Active Directory (Windows Server 2022)\n- 2 serveurs de fichiers (NAS Synology, 40 To)\n- 1 serveur de messagerie (Exchange 2019)\n- 1 serveur d'applications (Linux Ubuntu 22.04)\n- 2 serveurs de sauvegarde (Veeam Backup)"),
    ("Sécurité informatique", "Pare-feu : Fortinet FortiGate 200F\nAntivirus : Sophos Endpoint Protection (licences centralisées)\nSauvegardes : quotidiennes automatisées avec rétention 30 jours\nDernière mise à jour de sécurité : 28 mars 2025"),
    ("Projets IT 2025", "- Migration vers Microsoft 365 Cloud (Q2 2025) — Budget : 18 000 €\n- Déploiement de la signature électronique (Q3 2025) — Budget : 8 500 €\n- Remplacement de 20 postes vétustes (Q4 2025) — Budget : 28 000 €"),
])

make_csv(os.path.join(BASE, "IT", "tickets_support_2025.csv"),
    ["N_ticket", "Date", "Agent", "Service", "Catégorie", "Description", "Priorité", "Statut", "Résolution"],
    [
        ["TK-001", "2025-01-06", "M006 Thomas Lucas", "RH", "Matériel", "Écran ne s'allume plus", "Haute", "Résolu", "Remplacement écran"],
        ["TK-002", "2025-01-14", "M002 Lefebvre Marc", "Finance", "Logiciel", "Impossible d'accéder à CIVIL-NET", "Haute", "Résolu", "Réinitialisation droits"],
        ["TK-003", "2025-01-21", "M005 Bernard Claire", "Communication", "Réseau", "Connexion WiFi instable bureau 12", "Moyenne", "Résolu", "Remplacement borne WiFi"],
        ["TK-004", "2025-02-03", "M009 Laurent Nadia", "Comptabilité", "Logiciel", "Erreur lors de l'export Excel", "Basse", "Résolu", "Mise à jour Office"],
        ["TK-005", "2025-02-17", "M001 Durand Sophie", "Direction", "Sécurité", "Email suspect reçu", "Haute", "Résolu", "Sensibilisation phishing"],
        ["TK-006", "2025-03-05", "M007 Richard Émilie", "Achats", "Matériel", "Imprimante en panne", "Moyenne", "En cours", "Commande pièce en attente"],
        ["TK-007", "2025-03-19", "M004 Petit Antoine", "Urbanisme", "Logiciel", "GEO-MAIRIE lent au démarrage", "Basse", "En cours", "Diagnostic en cours"],
    ]
)

# ─── DIRECTION ───
make_docx(os.path.join(BASE, "Direction", "projet_de_ville_2025_2030.docx"), "Projet de Ville 2025–2030 — Villeneuve-les-Prés", [
    ("Vision stratégique", "La Mairie de Villeneuve-les-Prés s'engage dans une trajectoire de développement durable, inclusif et numérique pour les cinq prochaines années. Ce projet de ville a été co-construit avec les habitants lors de 6 réunions de concertation organisées en 2024."),
    ("Axe 1 — Transition écologique", "Objectifs : réduire de 30% les émissions de CO2 des bâtiments municipaux d'ici 2030, développer les mobilités douces (3 km de pistes cyclables supplémentaires), créer un jardin partagé de 2 000 m², et atteindre 60% d'énergies renouvelables dans la consommation municipale."),
    ("Axe 2 — Services à la population", "Modernisation de l'état civil avec la dématérialisation complète des démarches, extension des horaires de la médiathèque, création d'une Maison France Services, et déploiement d'une application mobile municipale."),
    ("Axe 3 — Développement économique", "Soutien à la création d'entreprises via la pépinière d'entreprises municipale, revitalisation du centre-bourg avec 8 nouvelles boutiques, et développement du marché hebdomadaire avec 15 nouveaux stands de producteurs locaux."),
    ("Axe 4 — Cohésion sociale", "Construction de 120 logements sociaux, développement de l'offre périscolaire, création d'un pôle santé regroupant 6 médecins généralistes, et programme d'inclusion numérique pour les seniors."),
    ("Gouvernance et suivi", "Un comité de pilotage trimestriel réunissant élus, services et représentants de la société civile assurera le suivi du projet. Un rapport annuel public sera présenté au conseil municipal chaque mois d'avril."),
])

make_docx(os.path.join(BASE, "Direction", "compte_rendu_conseil_municipal_mars2025.docx"), "Compte Rendu — Conseil Municipal du 18 mars 2025", [
    ("Présents", "Mme Fontaine Marie-Hélène (Maire), M. Garnier Jean-Pierre (1er adjoint), Mme Rousseau Béatrice (2e adjoint), M. Chevallier Alain (3e adjoint), et 14 conseillers municipaux. Excusés : M. Bonnet Frédéric, Mme Collet Sandrine."),
    ("Délibération n°1 — Budget primitif 2025", "Le conseil municipal approuve à l'unanimité le budget primitif 2025 s'élevant à 8 900 000 euros en fonctionnement et 1 200 000 euros en investissement. Vote : 17 pour, 0 contre, 1 abstention."),
    ("Délibération n°2 — Marché travaux voirie", "Attribution du marché de réfection de voirie à l'entreprise BTP Île-de-France pour un montant de 420 000 € HT. Vote : 15 pour, 2 contre, 1 abstention."),
    ("Délibération n°3 — Convention CCAS", "Renouvellement de la convention de partenariat avec le CCAS pour l'exercice 2025 avec une subvention de 95 000 euros. Vote : 18 pour, 0 contre, 0 abstention."),
    ("Questions diverses", "M. Chevallier interpelle la Maire sur l'avancement du projet de piste cyclable. Mme Fontaine confirme que les travaux débuteront en septembre 2025 et rappelle que la concertation publique est ouverte jusqu'au 30 avril 2025."),
])

# ─── MARKETING / COMMUNICATION ───
make_docx(os.path.join(BASE, "Marketing", "plan_communication_2025.docx"), "Plan de Communication 2025 — Mairie de Villeneuve-les-Prés", [
    ("Objectifs de communication", "Renforcer le lien entre la mairie et les habitants, valoriser les actions municipales, et promouvoir les événements et services de la collectivité. Cible prioritaire : les 25-45 ans et les familles."),
    ("Supports de communication", "- Bulletin municipal trimestriel (tirage : 4 500 exemplaires)\n- Site internet municipal (30 000 visiteurs/mois)\n- Réseaux sociaux : Facebook (2 800 abonnés), Instagram (1 200 abonnés)\n- Newsletter mensuelle (1 800 abonnés)\n- Panneaux d'affichage municipaux (12 emplacements)"),
    ("Calendrier éditorial T1-T2 2025", "- Janvier : Vœux du Maire, résultats du budget participatif\n- Février : Carnaval municipal, inscriptions périscolaires\n- Mars : Semaine du développement durable, marché de printemps\n- Avril : Compte rendu conseil municipal, lancement appli mobile\n- Mai : Fête des voisins, journée portes ouvertes services techniques\n- Juin : Fête de la musique, bilan mi-année"),
    ("Budget communication 2025", "Budget total alloué : 35 000 euros\n- Impression bulletin municipal : 12 000 €\n- Refonte site internet : 8 000 €\n- Réseaux sociaux et community management : 6 000 €\n- Événementiel et signalétique : 9 000 €"),
])

# ─── ACHATS ───
make_xlsx(os.path.join(BASE, "Achats", "plan_achats_2025.xlsx"), "Plan Achats 2025",
    ["N_ref", "Objet", "Service_demandeur", "Budget_prévisionnel", "Procédure", "Date_lancement", "Titulaire_prévu", "Statut"],
    [
        ["PA-001", "Fournitures de bureau", "Tous services", 45000, "MAPA", "2025-01-01", "OFFICE PLUS SAS", "Notifié"],
        ["PA-002", "Maintenance informatique", "IT", 78000, "MAPA", "2025-01-15", "TECHSOLUTIONS 95", "Notifié"],
        ["PA-003", "Nettoyage des locaux", "Patrimoine", 62000, "MAPA", "2025-02-01", "PROPRE & NET", "Notifié"],
        ["PA-004", "Travaux voirie", "Technique", 420000, "AO ouvert", "2025-02-15", "BTP ÎLE-DE-FRANCE", "Notifié"],
        ["PA-005", "Véhicule électrique de service", "Technique", 38000, "MAPA", "2025-04-01", "En cours", "Publication"],
        ["PA-006", "Matériel de sonorisation salle fêtes", "Culture", 12000, "MAPA", "2025-05-01", "À définir", "Préparation"],
        ["PA-007", "Logiciel de gestion des congés", "RH", 8500, "MAPA", "2025-06-01", "À définir", "Préparation"],
        ["PA-008", "Mobilier bureau open space", "Direction", 22000, "MAPA", "2025-07-01", "À définir", "Planifié"],
    ]
)

make_docx(os.path.join(BASE, "Achats", "politique_achats_responsables.docx"), "Politique d'Achats Responsables — Mairie de Villeneuve-les-Prés", [
    ("Engagement de la collectivité", "La Mairie de Villeneuve-les-Prés s'engage à intégrer des critères environnementaux, sociaux et éthiques dans l'ensemble de ses procédures d'achats, conformément aux dispositions du Code de la commande publique et du Plan National pour des Achats Durables (PNAD) 2022-2025."),
    ("Critères environnementaux", "Tout marché supérieur à 10 000 € HT doit intégrer au minimum une clause environnementale parmi les suivantes : produits labellisés (Écolabel européen, NF Environnement), véhicules à faibles émissions, matériaux recyclés ou recyclables, limitation des emballages."),
    ("Critères sociaux", "Les marchés de travaux et services supérieurs à 100 000 € HT doivent comporter une clause d'insertion professionnelle représentant au minimum 5% des heures travaillées. La mairie favorise également les ESAT et entreprises adaptées pour les achats de moins de 40 000 €."),
    ("Circuits courts et économie locale", "Pour les achats alimentaires (cantines, événements), la mairie vise 50% de produits locaux et bio d'ici 2026. Les TPE/PME locales sont encouragées à candidater via des lots adaptés à leur taille."),
    ("Suivi et évaluation", "Un tableau de bord annuel des achats responsables est présenté au conseil municipal. Indicateurs suivis : part des marchés avec clause environnementale, heures d'insertion réalisées, part des fournisseurs locaux."),
])

print("\n🎉 Tous les documents ont été créés avec succès !")